import os
import streamlit as st
import tiktoken
import pandas as pd

import io
from docx import Document

# TEST VERSION #
local_test = False
if local_test == True:
    from apikeys import openaikey

# CLOUD VERSION #
# openaikey = st.secrets['OPENAI_API_KEY']

# AI STUFF BELOW
# ____________________

#Import LLM - Langchain
from langchain.llms import OpenAI

#Langchain Features
from langchain.prompts import PromptTemplate
from langchain.chains import LLMChain

#############################################################

################### App Layout #########################

with st.sidebar:
    st.title("AI Survey Response Reports")
    st.header("Turn responses into insights")
    st.write("Upload Tally Forms survey results as a CSV & use AI to summarise all survey responses.")
    st.divider()
    user_openai = st.text_input("Please enter your OpenAI API Key:",)

report_name = st.text_input("Give a title to your survey response report.")

survey_responses_csv = st.file_uploader("Upload the survey response file here:", accept_multiple_files=False, type="csv")
submit = st.button("Submit")

if submit and report_name == "":
    st.error("You must enter an Organisation Name")
elif submit and report_name != "" and survey_responses_csv is None:
    st.error("You must upload the survey responses")
elif submit and local_test == False and user_openai == "":
    st.error("You must enter an OpenAI API Key")
elif submit and report_name != "" and survey_responses_csv is not None:
    if local_test == False:
        openaikey = user_openai
    os.environ['OPENAI_API_KEY'] = openaikey
    llm = OpenAI(temperature=0.4)

    response_summary_template = PromptTemplate(
        input_variables = ["survey_question", "responses"],
        template ='Give a summary of the thoughts and feelings expressed in the following survey responses, as an answer to the question {survey_question}. Highlight any recurring themes or particular areas of concern. Here are the survey responses: {responses}'
    )

    combined_summary_template = PromptTemplate(
        input_variables=["current_summary", "new_summary"],
        template = "Combine the following two summaries into one text. Give particular focus to any recurring feelings, themes and areas of concern. Summary 1: {current_summary}. Summary 2: {new_summary}"    
    )

    collective_feedback_template = PromptTemplate(
        input_variables = ["collective_feedback"],
        template = "Give an overall summary of the different ideas and opinions expressed in the following survey response summaries: {collective_feedback}. Begin the response with the phrase: 'Collective feedback from the survey shows that, overall, respondents feel that ...' "
    )

    ######## Theme Analysis Summary Chains ##########

    summary_chain = LLMChain(llm=llm, prompt = response_summary_template, verbose = False, output_key='summary')

    combined_summary_chain = LLMChain(llm=llm, prompt= combined_summary_template, verbose=False, output_key='combined_summary')

    collective_feedback_chain = LLMChain(llm=llm, prompt = collective_feedback_template, verbose = False, output_key='overall_summary')

    #_______________________________________
    #AI STUFF ABOVE

    ##################### Process Response Functions ###############
    total_token_count = 0

    def process_survey_responses(survey_question = str, responses = list):
        token_count = 0
        response_count = len(responses)
        processed_count = 0
        llm_call_buffer = []
        current_summary = ""
        new_summary = ""

        for response in responses:
            encoding = tiktoken.encoding_for_model("gpt-3.5-turbo")
            encoding = encoding.encode(response)
            token_count += len(encoding)
            llm_call_buffer.append(response)
            processed_count += 1

            if token_count <= 3000 and processed_count == response_count:
                new_summary = summary_chain.run({"survey_question": survey_question, "responses": llm_call_buffer})
            elif token_count > 3000 and processed_count < response_count:
                new_summary = summary_chain.run({"survey_question": survey_question, "responses": llm_call_buffer})
                token_count = 0
                llm_call_buffer = []
            
            if current_summary == "" and new_summary != "":
                current_summary = new_summary
                new_summary = ""

            if current_summary != "" and new_summary != "":
                combined_summary = collective_feedback_chain.run({"current_summary": current_summary, "new_summary": new_summary})
                current_summary = combined_summary
                token_count = 0
                llm_call_buffer = []
                new_summary = ""

        return current_summary

    with st.spinner("Processing Survey Data"):
        df = pd.read_csv(survey_responses_csv)
        responses = len(df.index)

        df_cols = len(df.axes[1])
        col_names = list(df.columns)

        for i in range(3):
            col_names.pop(0)    

        all_summaries = []
        num_responses = len(df.index)

        for col in col_names:
            col_responses = []
            for i in range(num_responses):
                individual_response = str(df.at[i, col])
                col_responses.append(individual_response)
            
            response_summary = process_survey_responses(col, col_responses)
            question_obj = {"Question": col, "Summary": response_summary}
            all_summaries.append(question_obj)

    st.success("Responses have been processed")

    st.subheader(f'Survey Diagnostic Report: {report_name}')

    for question_obj in all_summaries:
        i = (all_summaries.index(question_obj) + 1)
        question = question_obj["Question"]
        string_q = f"Question {i}: {question}"
        with st.expander(string_q):
            st.write(question_obj["Summary"])

    with st.spinner("Generating Report as Word Doc"):
        report = Document()

        paragraph_format = report.styles['Normal'].paragraph_format
        paragraph_format.space_after = None

        report.add_heading(f'Survey Diagnostic Report: {report_name}', level=1)

        for question_obj in all_summaries:
            i = (all_summaries.index(question_obj) + 1)
            question = question_obj["Question"]
            report.add_heading(f"Question {i}: {question}", level=3)
            report.add_paragraph(question_obj["Summary"])
            
        bio = io.BytesIO()
        report.save(bio)

        st.download_button(
            label="Download Report",
            data=bio.getvalue(),
            file_name=f"{report_name} Report.docx",
            mime="docx",
        )

        


